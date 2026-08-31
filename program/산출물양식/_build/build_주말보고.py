# -*- coding: utf-8 -*-
"""주말(2026-08-29 토 ~ 08-30 일) 작업 보고서를 만든다.

    python program/산출물양식/_build/build_주말보고.py

★내용은 지어내지 않는다. 커밋 로그와 리포트 파일에서만 가져온다.
  아래 COMMITS 는 `git log --since=... --until=...` 으로 뽑은 실제 목록이다.
  숫자(커밋 수, 증감 줄수)도 `--numstat` 실측이다.

★다른 산출물과 달리 이 문서는 채울 양식이 없다. 그래서 python-docx 로
  처음부터 만든다. `docxfill.py` 는 양식을 고치는 도구라 여기 쓸 수 없다.

★금지문자 검사는 그대로 건다. 줄표와 화살표는 이 저장소에서 안 쓴다.
"""
import os
import sys
from datetime import date

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FORMS = os.path.dirname(HERE)
OUT = os.path.join(FORMS, "[보고] 주말 작업 보고서_2026-08-29~30_A-COPilot.docx")

BANNED = "—–→⇒★☆✓✔✕▸…⇄①②③④⑤⑥"
INK = RGBColor(0x1F, 0x24, 0x30)
DIM = RGBColor(0x5C, 0x66, 0x7A)
ACCENT = RGBColor(0x2F, 0x5B, 0xD8)


def check(text):
    hit = [c for c in BANNED if c in text]
    if hit:
        raise ValueError("금지문자 %s 가 있다: %s" % (hit, text[:60]))
    return text


def font(run, size=10.5, bold=False, color=INK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def para(doc, text="", size=10.5, bold=False, color=INK, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    if text:
        font(p.add_run(check(text)), size, bold, color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run(check(text)), 14 if level == 1 else 11.5, True,
         ACCENT if level == 1 else INK)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(16)
    p.paragraph_format.line_spacing = 1.4
    font(p.add_run("- " + check(text)), size)
    return p


def table(doc, head, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(head):
        cell = t.rows[0].cells[i]
        cell.text = ""
        font(cell.paragraphs[0].add_run(check(h)), 9.5, True, DIM)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            font(cells[i].paragraphs[0].add_run(check(str(v))), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── 실측값 (git log --numstat, 2026-08-29 00:00 ~ 08-31 00:00) ───────────────
STATS = [
    ("루트 저장소", "26", "326", "+19,754 / -595"),
    ("final_project_sample", "12", "47", "+2,426 / -126"),
    ("합계", "38", "373", "+22,180 / -721"),
]

DEFECTS = [
    ("평가 채점기", "승인 대기를 실패로 세고 성공 판정을 뒤에서 덮어썼다",
     "A군 정확도가 0.0퍼센트로 나왔다", "지시서로 담당에게 넘김"),
    ("프롬프트 키 등록", "CS Pack 신규 키를 등록하지 않고 목록을 비워 뒀다",
     "Response Review 팀이 감사 경로로 불릴 때마다 죽었다", "수정 완료"),
    ("모듈 토글", "mcp 와 voc 에 게이트가 없고 graph_store 는 우회 경로가 있었다",
     "화면에서 꺼도 동작이 그대로였다", "수정 완료, 회귀 테스트 9건"),
    ("도표 잘림", "상자 폭을 손으로 적어 마지막 칸이 축 밖으로 나갔다",
     "Composer 흐름과 데이터 흐름 두 그림에서 마지막 칸이 안 보였다", "수정 완료"),
]

REMAIN = [
    ("택배 제출본 합치기", "raw/_incoming_20260829 를 processed 로",
     "네이버 4개는 형식이 같고 쿠팡 4개는 변환 필요"),
    ("네이버 4건 누락", "크롤러를 주문 링크 단위 순회로",
     "cyw 분에서만 확인됐고 나머지 3명분은 대조 안 함"),
    ("VOC 전처리", "4종 924MB", "aihub 3종과 kaggle 묶음"),
    ("평가 채점기 수정", "지시서 이행", "이게 끝나야 평가 수치를 믿을 수 있다"),
    ("파인튜닝 데이터", "response_review 를 켜서 운영 기록을 쌓기",
     "지금은 학습 데이터가 16건이 상한이다"),
]

COMMITS_SAT = [
    "화면설계서에 Playwright 실제 캡처와 목업 3종 추가",
    "구축 대행을 배제하는 서술과 대립 비교표 제거",
    "TeamFlow 스프린트 4개 생성과 에픽 13건 배정",
    "Composer v3 토글 엔드포인트와 acop_composer_ui 패키지 신설",
    "cs 에 introspection, outbox, 카탈로그 검증 Team 반영",
    "카탈로그 기반 인스턴스 만들기와 지우기 화면 연결",
    "중앙 설정 저장소 결정과 ConfigStore, AuditStore 계층",
    "시스템 구성도 작성, 평가 harness 결함 2건 기록과 지시서",
    "VOC 데이터셋 3종 리포트",
]
COMMITS_SUN = [
    "DoD-28 파인튜닝 파이프라인 완주와 RAG 통합 배선",
    "프롬프트 키 등록 결함 수정",
    "대상이 중앙 저장소에서 자기 선언을 읽는 경로",
    "감사 기록도 중앙으로, 설정 서비스 진입점 신설",
    "direct 와 central 두 운영 방식을 클라이언트와 UI 가 지원",
    "화면설계서에 설명만 있고 그림이 없던 카드 3장 추가",
    "도표 잘림 수정",
    "중간 발표회 발표자료 19장 작성",
    "모듈 토글 실효화와 회귀 테스트",
]

doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Pt(60)
sec.top_margin = sec.bottom_margin = Pt(50)

# ── 표지 머리 ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
font(p.add_run("SK 네트웍스 Family AI 32기 : 6팀 A-COPilot"), 9.5, False, DIM)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
font(p.add_run("주말 작업 보고서"), 20, True, INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
font(p.add_run("대상 기간 2026년 8월 29일 토요일부터 8월 30일 일요일까지  ·  "
               "작성 %s" % date.today().isoformat()), 9.5, False, DIM)

# ── 1. 한 줄 요약 ────────────────────────────────────────────────────────────
heading(doc, "1. 한 줄 요약")
para(doc, "주말 이틀 동안 커밋 38개를 남겼다. 가장 큰 줄기는 설정을 중앙 저장소로 옮기는 "
          "작업이었고, 두 번째는 파인튜닝 파이프라인을 끝까지 돌린 것이다. "
          "파인튜닝은 완주했지만 성능이 더 나빠 채택하지 않기로 판정했다.")
table(doc, ["저장소", "커밋", "파일 변경", "줄 증감"], STATS)

# ── 2. 무엇을 했나 ───────────────────────────────────────────────────────────
heading(doc, "2. 무엇을 했나")

heading(doc, "2.1 중앙 설정 저장소 (주말 최대 작업)", 2)
para(doc, "설정을 대상 서버의 파일이 아니라 우리가 운영하는 중앙 데이터베이스에 두기로 "
          "8월 29일에 결정하고 일요일까지 구현했다.")
para(doc, "이유는 하나다. 구성을 바꾸는 기능이 고객 대면 경로에 있으면 취약점 하나가 "
          "서비스 전체에 닿는다. 그래서 그 기능을 중앙 한 곳에만 두고, 고객사에 배포되는 "
          "대상은 기동할 때 자기 선언을 읽기만 하게 했다.", size=10)
para(doc, "만든 것", bold=True, space_after=3)
for t in ["설정과 감사 기록을 파일과 중앙 DB 어느 쪽이든 같은 방식으로 다루는 저장 계층",
          "대상이 중앙에서 자기 선언을 읽는 경로. 못 읽으면 조용히 넘어가지 않고 그 자리에서 멈춘다",
          "중앙 한 곳에서 여러 대상을 관리하는 설정 서비스 진입점",
          "직접 연결 방식과 중앙 방식 두 가지를 클라이언트와 화면이 모두 지원",
          "마이그레이션 2개와 통합 테스트 5개"]:
    bullet(doc, t)

heading(doc, "2.2 Composer 카탈로그와 인스턴스 관리", 2)
para(doc, "운영자가 화면에서 모듈과 Team 을 고르고 켜고 끌 수 있게 하는 작업이다. "
          "카탈로그에서 골라 인스턴스를 만들고 지우는 화면까지 연결했다. "
          "화면이 쓰던 자체 전송 구현은 공용 패키지로 뺐다. 같은 코드가 두 곳에 있으면 "
          "한쪽만 고쳐지기 때문이다.")

heading(doc, "2.3 DoD-28 파인튜닝", 2)
para(doc, "GPU 서버에서 파이프라인을 끝까지 돌렸다. 결과는 채택 불가다.")
table(doc, ["비교", "Proposed", "Proposed 더하기 파인튜닝"],
      [["holdout 24건 통과율", "16.7퍼센트", "0.0퍼센트"],
       ["golden", "같은 방향", "같은 방향"]])
para(doc, "학습 지표는 좋아졌는데 실제 출력이 나빠졌다. 재비교에서도 손실은 1.65에서 "
          "1.48로, 정확도는 0.67에서 0.71로 개선됐지만 실제 출력은 유효한 형식을 "
          "하나도 만들지 못했다.", size=10)
para(doc, "원인은 확인됐다. 이 팀의 학습 데이터가 16건이 상한이다. 더 늘리려면 응답 검토 "
          "기능을 실제로 켜서 운영 기록을 쌓는 방법밖에 없다.", size=10)

heading(doc, "2.4 산출물 문서", 2)
for t in ["화면설계서에 실제 화면 캡처와 목업 3종 추가. 이후 설명만 있고 그림이 없던 카드 3장도 보강",
          "시스템 구성도 작성",
          "중간 발표회 발표자료 19장 작성",
          "구축 대행을 배제하던 서술과 대립 비교표 제거"]:
    bullet(doc, t)

heading(doc, "2.5 일정 관리", 2)
para(doc, "TeamFlow 에 스프린트 4개를 만들고 에픽 13건을 배정했다. "
          "권한 응답이 401에서 403으로 바뀐 것과, 이미 매겨진 이슈 키를 다시 매기지 "
          "않기로 한 결정을 함께 기록했다.")

# ── 3. 결함 ──────────────────────────────────────────────────────────────────
heading(doc, "3. 찾아서 고친 결함")
para(doc, "네 건이다. 세 건은 고쳤고 한 건은 담당 영역이라 지시서로 넘겼다.", size=10)
table(doc, ["결함", "무엇이 잘못돼 있었나", "증상", "처리"], DEFECTS)

# ── 4. 숨기지 않는 것 ────────────────────────────────────────────────────────
heading(doc, "4. 숨기지 않고 적는 것")
para(doc, "파인튜닝은 실패가 아니라 측정해서 안 쓰기로 판정한 것이다.", bold=True)
para(doc, "돌려 보지 않고 넘어갔다면 이 결론을 낼 수 없었다. 다만 발표에서 "
          "성과로 말하면 안 된다. 채택하지 않은 이유와 데이터가 모자란다는 사실을 "
          "함께 말해야 한다.", size=10)
para(doc, "평가 수치는 아직 믿을 수 없다.", bold=True)
para(doc, "채점기 결함이 아직 안 고쳐졌다. 이 수정이 끝나기 전의 정확도 수치는 "
          "그대로 쓰면 안 된다.", size=10)

# ── 5. 남은 일 ───────────────────────────────────────────────────────────────
heading(doc, "5. 남은 일")
table(doc, ["항목", "무엇을", "메모"], REMAIN)

# ── 6. 근거 ──────────────────────────────────────────────────────────────────
heading(doc, "6. 근거")
para(doc, "아래는 커밋 제목을 사람이 읽을 수 있게 줄인 것이다. "
          "원본은 git log 로 확인한다.", size=10, color=DIM)
para(doc, "8월 29일 토요일", bold=True, space_after=3)
for t in COMMITS_SAT:
    bullet(doc, t, size=10)
para(doc, "")
para(doc, "8월 30일 일요일", bold=True, space_after=3)
for t in COMMITS_SUN:
    bullet(doc, t, size=10)
para(doc, "")
para(doc, "재현 명령", bold=True, space_after=3)
para(doc, 'git log --since="2026-08-29 00:00" --until="2026-08-31 00:00" '
          '--pretty="%ad %h %s"', size=9.5, color=DIM)
para(doc, "이 문서는 program/산출물양식/_build/build_주말보고.py 가 만든다. "
          "손으로 고치지 않는다.", size=9, color=DIM)

doc.save(OUT)
print("저장:", OUT)
